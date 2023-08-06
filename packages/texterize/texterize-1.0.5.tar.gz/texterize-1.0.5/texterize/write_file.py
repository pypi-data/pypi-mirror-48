import os, numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor

FONT = "Courier"

#helper function to 
def write(text_arr, chroma, output_file_type, write_path, overwrite):
    if output_file_type == "HTML":
        writeHTML(text_arr, chroma, write_path, overwrite)
    elif output_file_type == "Word":
        writeDoc(text_arr, chroma, write_path, overwrite)
    else:
        raise Exception("Error writing file.")

#writes a text block to a .HTML file
def writeHTML(text_arr, chroma, write_path, overwrite):
    '''
    params:
    -text_arr: A 2-D numpy array of the text to write to the document
    -chroma: A 3-D numpy array containing the RGB values to assign to each character of text
    -write_path: The path to write the output file to
    -overwrite: A Boolean representing whether or not to overwrite an existing .docx file if found.

    return:
    -the path to the output file
    '''
    FONT_SIZE = 750 / min(text_arr.shape[0], text_arr.shape[1])
    LINE_SPACING = FONT_SIZE * 0.45

    try:
        if not write_path[:-5] == ".html":
            write_path += ".html"
        if overwrite and os.path.exists(write_path):
            os.remove(write_path)
    except:
        raise Exception(f"Invalid write path {write_path}")
    
    doc = open(write_path, "w")
    doc.write(f'''
    <html>
        <head>
            <title>Texterize</title>
        </head>
        <body style='font-size:{FONT_SIZE}pt; line-height:{LINE_SPACING}pt; font-family:courier; letter-spacing:-1;'>
    ''')

    for i in range(text_arr.shape[0]):
        doc.write(f"<div>")
        for j in range(text_arr.shape[1]):
            R, G, B = [int(c) for c in chroma[i][j]]
            doc.write(f"<text style='color:rgb({R},{G},{B});'>{text_arr[i][j]}</text>")
        doc.write("</div>")

    doc.write('''
        </body>
    </html>
    ''')
    doc.close()

    return write_path

#writes a text block to a .docx file
def writeDoc(text_arr, chroma, write_path, overwrite):
    '''
    params:
    -text_arr: A 2-D numpy array of the text to write to the document
    -chroma: A 3-D numpy array containing the RGB values to assign to each character of text
    -write_path: The path to write the output file to
    -overwrite: A Boolean representing whether or not to overwrite an existing .docx file if found.

    return:
    -the path to the output file
    '''
    FONT_SIZE = 400 / min(text_arr.shape[0], text_arr.shape[1])
    LINE_SPACING = 0.65

    document = Document()
    doc_style = document.styles["Normal"]

    font = doc_style.font
    font.name = FONT
    font.size = Pt(FONT_SIZE)

    p_format = doc_style.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_format.space_before = 0
    p_format.space_after = 0
    p_format.line_spacing = LINE_SPACING

    for i in range(text_arr.shape[0]):
        p = document.add_paragraph()
        p.style = doc_style
        for j in range(text_arr.shape[1]):
            c = p.add_run(text_arr[i][j])
            R, G, B = [int(c) for c in chroma[i][j]]
            c.font.color.rgb = RGBColor(R, G, B)

    if not write_path[:-5] == ".docx" and not write_path[:-5] == ".doc":
        write_path += ".docx"
    try:
        if overwrite and os.path.exists(write_path):
            os.remove(write_path)
        document.save(write_path)
    except:
        raise Exception(f"Invalid write path {write_path}")

    return write_path